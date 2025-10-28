import docx
import pytesseract
import re
import io
from PIL import Image
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Path to tesseract
pytesseract.pytesseract.tesseract_cmd = r"C:/Program Files/Tesseract-OCR/tesseract.exe"

# ---------------------
# Cleaning helper
# ---------------------
def clean_text(text: str) -> str:
    text = re.sub(r'\n+', '\n', text)          # collapse multiple newlines
    text = re.sub(r'[ \t]+', ' ', text)        # collapse multiple spaces
    text = re.sub(r' *\n *', '\n', text)       # trim spaces around newlines
    return text.strip()

# ---------------------
# Extract text & images
# ---------------------
def extract_text_from_docx(docx_path: str) -> str:
    final_content = []

    try:
        doc = docx.Document(docx_path)

        # --- 1. Extract paragraphs (main body text) ---
        for para in doc.paragraphs:
            if para.text.strip():
                final_content.append(para.text.strip())

        # --- 2. Extract tables (cell-by-cell text) ---
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    final_content.append(row_text)

        # --- 3. Extract images and OCR them ---
        # DOCX stores images as relationships (rId → binary image)
        rels = doc.part.rels
        for rel in rels.values():
            if rel.reltype == RT.IMAGE:
                image_data = rel.target_part.blob
                try:
                    pil_img = Image.open(io.BytesIO(image_data))
                    ocr_text = pytesseract.image_to_string(pil_img).strip()
                    if ocr_text:
                        final_content.append(ocr_text)
                except Exception as e:
                    print(f"[Warning] OCR failed on image: {e}")

    except Exception as e:
        print(f"[Critical Error] Could not open DOCX: {e}")

    # --- Final clean ---
    combined = "\n".join(final_content)
    return clean_text(combined)

# ---------------------
# Example usage
# ---------------------
if __name__ == "__main__":
    docx_path = "./Keyword Extraction/Sample.docx"
    extracted_text = extract_text_from_docx(docx_path)

    print("Final Extracted Text:\n")
    print(extracted_text)
